from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            cells[i].text = str(cell_text)
    # Espacio después de la tabla
    doc.add_paragraph('')


def add_key_values(doc: Document, pairs):
    for k, v in pairs:
        p = doc.add_paragraph()
        run_k = p.add_run(f"{k}: ")
        run_k.bold = True
        p.add_run(v)


def build_document() -> Document:
    doc = Document()

    # Metadata
    today = datetime.now().strftime('%Y-%m-%d')
    doc.core_properties.title = 'Descripción Técnica General del Sistema Alerta Raven'
    doc.core_properties.author = 'Equipo Arquitectura - Alerta Raven'
    doc.core_properties.created = datetime.now()

    # Title
    add_title(doc, 'Descripción Técnica General del Sistema Alerta Raven')
    add_paragraph(doc, '')

    add_key_values(doc, [
        ('Versión', '1.0'),
        ('Fecha', today),
        ('Autor', 'Equipo Arquitectura - Alerta Raven'),
        ('Estado', 'Aprobado para difusión interna'),
    ])

    # Resumen Ejecutivo
    add_heading(doc, 'Resumen Ejecutivo', level=1)
    add_paragraph(doc, (
        'Alerta Raven es un sistema de monitoreo y gestión de incidentes compuesto por '
        'una API backend (FastAPI), un frontend web tipo dashboard (PWA) y una aplicación '
        'móvil cliente. Permite recibir, procesar y visualizar alertas en tiempo real, '
        'con autenticación básica y capacidades opcionales de clasificación por ML.'
    ))

    # Arquitectura General
    add_heading(doc, 'Arquitectura General', level=1)
    add_bullets(doc, [
        'Capa de presentación (web): HTML fragmentado en componentes; JavaScript para navegación y mapas.',
        'Capa de API: FastAPI con endpoints, autenticación JWT, tareas en background y WebSockets.',
        'Capa de datos: SQLite asíncrono (aiosqlite) con modelos y CRUD.',
        'Capa de inteligencia: modelo ML opcional y heurísticas para detección de eventos.'
    ])

    # Backend API
    add_heading(doc, 'API Backend (FastAPI)', level=1)
    add_key_values(doc, [
        ('Ubicación', 'c:\\Users\\Alejandro\\Desktop\\AlertaRaven5\\api\\main.py'),
        ('Tecnología', 'Python, FastAPI, Uvicorn, aiosqlite, WebSockets'),
    ])
    add_heading(doc, 'Funcionalidad clave', level=2)
    add_bullets(doc, [
        'Autenticación web: POST /api/auth/login, cookie JWT httponly; GET /logout.',
        'Protección de dashboard: verificación de cookie (verify_web_auth).',
        'Salud y utilidades: GET /health; montaje de estáticos en /static.',
        'Tiempo real: WebSocket manager para notificar eventos a clientes.',
        'Tareas de fondo: latido (heartbeat), snapshots de métricas y carga de modelo ML.',
        'Persistencia: SQLite (alertas.db) con acceso asíncrono.'
    ])
    add_heading(doc, 'Catálogo de Endpoints', level=2)
    add_paragraph(doc, 'Listado resumido de rutas principales con método, propósito y autenticación:')
    add_table(doc,
              headers=['Método', 'Ruta', 'Propósito', 'Auth'],
              rows=[
                  ['GET', '/', 'Raíz y estado básico', '—'],
                  ['GET', '/health', 'Verificación de salud', '—'],
                  ['GET', '/manifest.webmanifest', 'Manifest PWA', '—'],
                  ['GET', '/service-worker.js', 'Service Worker PWA', '—'],
                  ['WEBSOCKET', '/ws', 'Evento tiempo real dashboard', 'Cookie JWT'],
                  ['GET', '/dashboard', 'Dashboard protegido', 'Cookie JWT'],
                  ['GET', '/login', 'Página de acceso', '—'],
                  ['POST', '/api/auth/login', 'Login y emisión cookie', 'Credenciales admin'],
                  ['GET', '/logout', 'Cierre de sesión', 'Cookie JWT'],
                  ['POST', '/api/v1/emergency-alert', 'Recepción alerta móvil', 'API Key'],
                  ['POST', '/api/v1/emergency-alert-debug', 'Recepción alerta (debug)', '—'],
                  ['GET', '/api/v1/alerts/{alert_id}', 'Estado de alerta por ID', '—'],
                  ['GET', '/api/v1/alerts', 'Listado de alertas', '—'],
                  ['GET', '/api/alerts', 'Listado (vista dashboard)', 'Cookie JWT'],
                  ['GET', '/api/alerts/{alert_id}', 'Detalle de alerta (dashboard)', 'Cookie JWT'],
                  ['PUT', '/api/alerts/{alert_id}/status', 'Actualizar estado de alerta', 'Cookie JWT'],
                  ['POST', '/api/v1/sensor-events', 'Ingesta eventos de sensor', 'API Key'],
                  ['GET', '/api/v1/sensor-events', 'Listar eventos de sensor', 'API Key'],
                  ['GET', '/api/v1/sensor-events/export', 'Exportar CSV eventos', 'API Key'],
                  ['GET', '/api/statistics', 'Estadísticas para dashboard', 'Cookie JWT'],
                  ['GET', '/api/v1/metrics/sensor-events-summary', 'Resumen de sensores', '—'],
                  ['GET', '/api/v1/metrics/model', 'Métricas del modelo', '—'],
                  ['POST', '/api/v1/metrics/model/snapshot', 'Snapshot métricas', '—'],
                  ['POST', '/api/v1/model/train/randomforest', 'Entrenar modelo RF', 'API Key'],
                  ['POST', '/api/v1/model/predict', 'Predicción del modelo', '—'],
                  ['GET', '/api/v1/model/status', 'Estado del modelo', '—'],
                  ['GET', '/api/v1/metrics/model/history', 'Histórico de métricas', '—'],
                  ['GET', '/api/v1/contacts/{device_id}', 'Obtener contactos por dispositivo', 'API Key'],
                  ['PUT', '/api/v1/contacts/{device_id}', 'Reemplazar contactos', 'API Key'],
                  ['GET', '/api/v1/model/config', 'Config del modelo', '—'],
                  ['POST', '/api/v1/model/config', 'Actualizar config modelo', '—'],
                  ['GET', '/api/v1/model/dataset/summary', 'Resumen dataset entrenamiento', '—'],
              ])
    add_heading(doc, 'Integraciones y configuración', level=2)
    add_bullets(doc, [
        'Base de datos: aiosqlite sobre alertas.db.',
        'WebSockets: broadcasting de alertas y cambios de estado.',
        'ML: RFAccidentClassifier (ml.py) con umbral ACCIDENT_CONFIDENCE_THRESHOLD.',
        'Variables entorno: ALERTARAVEN_WEB_SECRET, ALERTARAVEN_ADMIN_USER/PASS, '
        'METRICS_SNAPSHOT_INTERVAL, ML_MODEL_DIR/ML_MODEL_PATH/ML_META_PATH.'
    ])

    # Frontend Web
    add_heading(doc, 'Frontend Web (Dashboard PWA)', level=1)
    add_key_values(doc, [
        ('Ubicación', 'c:\\Users\\Alejandro\\Desktop\\AlertaRaven5\\api\\static\\'),
        ('Tecnología', 'HTML, CSS, JavaScript (Leaflet, PWA)'),
    ])
    add_heading(doc, 'Estructura y componentes', level=2)
    add_bullets(doc, [
        'Entradas: index.html (dashboard), login.html (acceso), landing.html (pública).',
        'Componentes: static/components/*.html (dashboard, map, alerts, statistics, system).',
        'Loader: components-loader.js para inyección dinámica y evento componentsLoaded.',
        'Lógica: app.js para navegación, filtros, mapas, tablas y eventos de UI.',
        'PWA: manifest.webmanifest y service-worker.js para caché y rendimiento.',
        'Mapas: Leaflet con invalidación de tamaño en reflow y resize.'
    ])

    # Aplicación móvil
    add_heading(doc, 'Aplicación Móvil (Cliente)', level=1)
    add_bullets(doc, [
        'Función: detectar eventos de accidente/incidente y enviar alertas a la API.',
        'Datos: tipo de incidente, confianza, timestamp y ubicación.',
        'Flujo: petición HTTP segura; la API persiste y procesa en background.',
        'Integración: vía endpoints públicos expuestos por la API.'
    ])
    add_heading(doc, 'Autenticación de la App Móvil', level=2)
    add_bullets(doc, [
        'Esquema HTTP Bearer con clave de API (HTTPBearer + verify_api_key).',
        'Endpoints protegidos: /api/v1/sensor-events*, /api/v1/model/train/randomforest, contactos.',
        'Rotación de claves recomendada y almacenamiento seguro en el cliente.'
    ])

    # Modelo de Datos
    add_heading(doc, 'Modelo de Datos', level=1)
    add_bullets(doc, [
        'Entidades: EmergencyAlert, AlertStatus, AccidentType, DeviceInfo, NotificationLog, SensorEvent, SensorEventType.',
        'Estados de alerta: received, processing, confirmed/completed, pending_review, failed.'
    ])
    add_heading(doc, 'Modelos Pydantic Clave', level=2)
    add_bullets(doc, [
        'LocationData, MedicalInfo, EmergencyContact.',
        'AccidentEventData y EmergencyAlertRequest (payloads de alertas).',
        'AlertResponse (respuesta estándar de alertas).',
        'SensorEventIn (ingesta de eventos de sensor).',
        'TrainParams, PredictIn, ModelConfigIn, DeviceContactsResponse.'
    ])

    # Tiempo Real
    add_heading(doc, 'Tiempo Real (WebSockets)', level=1)
    add_bullets(doc, [
        'Eventos: nueva alerta, cambio de estado, métricas y heartbeats.',
        'Cliente: el dashboard actualiza tablas, badges y capas del mapa en tiempo real.'
    ])

    # Autenticación y seguridad
    add_heading(doc, 'Autenticación y Seguridad', level=1)
    add_bullets(doc, [
        'Actual: login con cookie JWT httponly, samesite=lax; rutas protegidas verifican token.',
        'Mejora propuesta: 2FA (TOTP/WebAuthn) previo a emitir JWT; cambios en main.py y static/login.html.',
        'CORS: permisivo en desarrollo; restringir dominios en producción.',
        'Producción: usar HTTPS y cookies secure=True.'
    ])
    add_heading(doc, 'Flujos de Sesión', level=2)
    add_bullets(doc, [
        'Web: credenciales admin → emisión JWT → acceso a dashboard y APIs web.',
        'Móvil: clave API en Authorization: Bearer → acceso a endpoints protegidos.',
        'Logout: elimina cookie y redirige a /landing.'
    ])
    add_heading(doc, 'Manejo de Errores y Logs', level=2)
    add_bullets(doc, [
        'Uso de HTTPException con códigos apropiados (401, 404, 500, 503).',
        'Health check retorna 503 ante fallos; login inválido responde 401 JSON.',
        'Logging informativo en operaciones clave (guardado de alertas, entrenamiento, exportaciones).'
    ])

    # Operación y despliegue
    add_heading(doc, 'Operación y Despliegue', level=1)
    add_bullets(doc, [
        'Desarrollo: python -m uvicorn main:app --host 127.0.0.1 --port 8000 --reload.',
        'Acceso: http://127.0.0.1:8000/login y dashboard protegido por cookie.',
        'Estáticos: servidos bajo /static.',
        'Base de datos: SQLite (alertas.db) inicializada en lifespan.',
        'Salud: endpoint /health para verificaciones durante despliegue.'
    ])
    add_heading(doc, 'Consideraciones de Rendimiento y Escalabilidad', level=2)
    add_bullets(doc, [
        'Índices en tablas críticas (alerts, sensor_events, snapshots) para consultas rápidas.',
        'Tareas en background para operaciones pesadas (procesamiento de alertas, heartbeats).',
        'Posible migración a Postgres y WebSocket gateway gestionado si aumenta la carga.',
        'Cache y PWA para mejorar tiempos de carga del frontend.'
    ])

    # Flujo extremo a extremo
    add_heading(doc, 'Flujo Extremo a Extremo', level=1)
    add_bullets(doc, [
        'La app móvil detecta evento y envía alerta a la API.',
        'La API valida, persiste, procesa en background y notifica por WebSocket.',
        'El dashboard recibe eventos y actualiza UI (tabla y mapa).',
        'El usuario filtra, consulta detalle y gestiona estados según permisos.'
    ])

    # Rutas y archivos clave
    add_heading(doc, 'Rutas y Archivos Clave', level=1)
    add_bullets(doc, [
        'API: main.py, database.py, models.py, ml.py, websocket_manager.py.',
        'Frontend: static/index.html, static/login.html, static/app.js, static/components-loader.js, static/styles.css.',
        'Componentes: static/components/*.html (dashboard, map, alerts, statistics, system).',
        'Datos: alertas.db y data/alertas.db (copia).',
        'Pruebas: api/tests/.'
    ])
    add_heading(doc, 'Pruebas y Cobertura', level=2)
    add_bullets(doc, [
        'Unitarias: lógica de base de datos, ML y utilidades.',
        'Integración: endpoints principales (alertas, sensor-events, estadísticas).',
        'WebSocket: conexión y broadcasting.',
        'Rendimiento: tiempos de respuesta y exportaciones.'
    ])

    # Buenas prácticas
    add_heading(doc, 'Buenas Prácticas y Consideraciones', level=1)
    add_bullets(doc, [
        'Seguridad: mantener secretos JWT en variables de entorno; usar HTTPS y cookies seguras.',
        'Rendimiento: tareas pesadas en background; cuidado con reflow de mapas.',
        'Escalabilidad: migrar BD y WebSockets a servicios gestionados si la carga crece.',
        'Observabilidad: logs claros de estado y actividad; uso de /health.'
    ])

    # Anexos: Variables de entorno
    add_heading(doc, 'Anexo: Variables de Entorno Clave', level=1)
    add_bullets(doc, [
        'ALERTARAVEN_WEB_SECRET: secreto para firmar JWT del dashboard.',
        'ALERTARAVEN_ADMIN_USER / ALERTARAVEN_ADMIN_PASS: credenciales de acceso web.',
        'METRICS_SNAPSHOT_INTERVAL: intervalo para snapshot de métricas.',
        'ML_MODEL_DIR / ML_MODEL_PATH / ML_META_PATH: rutas de modelo ML y metadatos.',
        'ACCIDENT_CONFIDENCE_THRESHOLD: umbral de confianza de clasificación.'
    ])
    add_heading(doc, 'Anexo: Esquema de Base de Datos (Resumen)', level=1)
    add_bullets(doc, [
        'emergency_alerts(alert_id, device_id, user_id, accident_type, timestamp, confidence, accel/gyro magnitudes, location_data, medical_info, emergency_contacts, status, created_at, updated_at, additional_data).',
        'devices(device_id, user_id, device_model, os_version, app_version, last_seen, is_active).',
        'notification_logs(log_id, alert_id, notification_type, recipient, status, timestamp, error_message).',
        'device_contacts(device_id, name, phone, relationship, is_primary, updated_at).',
        'sensor_events(event_id, device_id, label, timestamp, metrics, predicted_label, prediction_confidence, raw_data, created_at).',
        'model_metrics_snapshots(id, created_at, model_version, precision, recall, f1, classes_json, confusion_matrix_json).'
    ])

    add_heading(doc, 'Roadmap y Mejoras Futuras', level=1)
    add_bullets(doc, [
        'Implementación de 2FA (TOTP/WebAuthn) en el flujo web.',
        'Optimización y despliegue de modelo ML más robusto con pipelines.',
        'Hardening de seguridad (CSP, HSTS, rotación de claves API, roles).',
        'Observabilidad avanzada (tracing, métricas de negocio, alertas operativas).'
    ])

    return doc


def main():
    doc = build_document()
    output_path = 'docs/Descripcion_Tecnica_Sistema_AlertaRaven.docx'
    doc.save(output_path)
    print(f'Documento generado: {output_path}')


if __name__ == '__main__':
    main()