import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    raise SystemExit("python-docx no está instalado. Ejecuta: python -m pip install python-docx")


def add_title(doc: Document, text: str):
    doc.add_paragraph(text).style = doc.styles['Title']


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def add_kv(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def build_document() -> Document:
    doc = Document()

    # Encabezado y metadatos
    add_title(doc, "DOCUMENTO DE SOLICITUD DE CAMBIO/PROYECTO")
    add_kv(doc, "Fecha", "13/11/2025")
    add_kv(doc, "De", "Jorge Galindo")
    add_kv(doc, "Para", "Consejo de Arquitectos")
    add_kv(doc, "Copia (CC)", "Jefes de Proyecto afectados, Product Owners, Seguridad")
    add_kv(doc, "Asunto", "Solicitud Formal para la Implementación de Autenticación de Dos Factores (2FA) en Alerta Raven — API módulo main.py y Dashboard web")

    add_paragraph(doc, "---")

    # 1. Resumen Ejecutivo
    add_heading(doc, "1. Resumen Ejecutivo", level=1)
    add_paragraph(doc,
                  "Se solicita formalmente la aprobación y los recursos para implementar Autenticación de Dos Factores (2FA) en el sistema Alerta Raven, enfocada en el módulo main.py de la API y el Dashboard web. 2FA añade una capa adicional de protección más allá de contraseñas tradicionales, reforzando la seguridad de accesos, protegiendo información sensible y cumpliendo estándares y normativas vigentes. La iniciativa es clave para mitigar riesgos asociados a vulneración de credenciales y accesos no autorizados.")

    # 2. Justificación de la Solicitud
    add_heading(doc, "2. Justificación de la Solicitud", level=1)
    add_bullets(doc, [
        "Riesgo de Compromiso de Credenciales: contraseñas vulnerables a phishing, keyloggers, fuerza bruta y reutilización.",
        "Protección de Información Sensible: datos de alertas, ubicación, información de usuarios y componentes del proyecto.",
        "Cumplimiento Normativo: alineación con LFPDPPP (México), ISO 27001 y mejores prácticas de seguridad.",
        "Reducción de Riesgos: 2FA bloquea la mayoría de accesos no autorizados incluso con credenciales comprometidas.",
        "Mejora de la Confianza: incrementa la confianza de usuarios, clientes y socios en la robustez del sistema.",
    ])

    # 3. Alcance de la Implementación
    add_heading(doc, "3. Alcance de la Implementación", level=1)
    add_bullets(doc, [
        "Sistemas Prioritarios: API (main.py) y Dashboard web protegidos por cookie JWT (access_token).",
        "Tipos de Usuarios: Administradores y usuarios con acceso a información sensible del Dashboard.",
        "Exclusiones (Fase Inicial): Usuarios básicos con acceso limitado; se propone expansión por fases futuras.",
    ])

    # 4. Opciones de Implementación de 2FA
    add_heading(doc, "4. Opciones de Implementación de 2FA (Recomendadas)", level=1)

    add_heading(doc, "4.1 Aplicaciones Autenticadoras (TOTP)", level=2)
    add_bullets(doc, [
        "Descripción: códigos temporales de un solo uso generados por apps (Google Authenticator, Authy).",
        "Ventajas: alto nivel de seguridad, sin costo por SMS, fácil adopción.",
        "Implementación: almacenar secreto TOTP por usuario; validar código en /api/auth/mfa/verify usando pyotp.",
    ])

    add_heading(doc, "4.2 Código por SMS/Email", level=2)
    add_bullets(doc, [
        "Descripción: envío de código de un solo uso por SMS o email.",
        "Ventajas: accesible para usuarios sin app autenticadora.",
        "Consideraciones: costo por SMS, seguridad del canal de correo; recomendable como respaldo.",
    ])

    add_heading(doc, "4.3 WebAuthn (Llaves de seguridad / biometría)", level=2)
    add_bullets(doc, [
        "Descripción: autenticación sin contraseña o como segundo factor usando llaves FIDO2/biometría.",
        "Ventajas: máxima seguridad y experiencia moderna.",
        "Consideraciones: mayor complejidad de implementación; viable en fases posteriores.",
    ])

    add_heading(doc, "4.4 Códigos de Respaldo (Backup Codes)", level=2)
    add_bullets(doc, [
        "Descripción: conjunto de códigos únicos para usos de emergencia.",
        "Uso: permitir recuperación cuando el usuario no tiene acceso a su segundo factor.",
    ])

    # 5. Plan de Implementación en Alerta Raven
    add_heading(doc, "5. Plan de Implementación en Alerta Raven (main.py y frontend)", level=1)
    add_bullets(doc, [
        "Backend (main.py): agregar endpoints /api/auth/mfa/challenge y /api/auth/mfa/verify.",
        "Emisión de JWT (cookie access_token) solo tras verificación exitosa en mfa/verify.",
        "Dependencia: pyotp en requirements.txt para TOTP; considerar librerías WebAuthn en fases futuras.",
        "Frontend (static/login.html): flujo en dos pasos (credenciales → código 2FA).",
        "Pruebas: unitarias de TOTP y de endpoints; integración del flujo completo y regresiones de login protectivo.",
    ])

    # 6. Impacto Esperado
    add_heading(doc, "6. Impacto Esperado", level=1)
    add_bullets(doc, [
        "Seguridad reforzada de accesos administrativos.",
        "Reducción de probabilidad de incidentes por credenciales comprometidas.",
        "Cumplimiento reforzado de buenas prácticas y normativas.",
    ])

    # 7. Riesgos y Mitigaciones
    add_heading(doc, "7. Riesgos y Mitigaciones", level=1)
    add_bullets(doc, [
        "Riesgo de bloqueo de usuarios: ofrecer códigos de respaldo y proceso de recuperación.",
        "Complejidad adicional: documentación clara y soporte al usuario.",
        "Despliegue: realizar pruebas en entorno de desarrollo y pilotos controlados antes de producción.",
    ])

    # 8. Cronograma Tentativo
    add_heading(doc, "8. Cronograma Tentativo", level=1)
    add_bullets(doc, [
        "Semana 1: Diseño y endpoints backend (challenge/verify) y actualización de requirements.txt.",
        "Semana 2: Integración frontend en login y pruebas unitarias/integración.",
        "Semana 3: Piloto con usuarios administrativos y ajustes; preparación de despliegue.",
    ])

    # 9. Costos y Recursos
    add_heading(doc, "9. Costos y Recursos", level=1)
    add_bullets(doc, [
        "Recursos humanos: desarrollador backend, QA, seguridad y apoyo frontend.",
        "Herramientas: pyotp (sin costo), opcional servicio de SMS (coste variable).",
    ])

    # 10. Requerimientos Técnicos
    add_heading(doc, "10. Requerimientos Técnicos", level=1)
    add_bullets(doc, [
        "Almacenamiento seguro de secretos TOTP por usuario.",
        "Gestión de sesión: cookie httponly (access_token) emitida tras 2FA.",
        "Registro y auditoría: logs de intentos 2FA (éxito/fallo).",
    ])

    # 11. Criterios de Éxito
    add_heading(doc, "11. Criterios de Éxito", level=1)
    add_bullets(doc, [
        "Inicio de sesión válido solo con segundo factor correcto.",
        "Cobertura de pruebas y validaciones manuales aprobadas.",
        "Sin regresiones en /login y /dashboard; documentación actualizada.",
    ])

    # 12. Solicitud de Aprobación
    add_heading(doc, "12. Solicitud de Aprobación", level=1)
    add_paragraph(doc,
                  "Se solicita la aprobación del Consejo de Arquitectos para proceder con el plan de implementación descrito. Tras la aprobación, se iniciará la ejecución conforme al cronograma y se informará al comité de avances y resultados.")

    return doc


def main():
    out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.abspath(os.path.join(out_dir, "Solicitud_Cambio_2FA_AlertaRaven.docx"))
    doc = build_document()
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()