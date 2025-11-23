import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    raise SystemExit("python-docx no está instalado. Ejecuta: python -m pip install python-docx")


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')


def add_kv_paragraph(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)


def build_document() -> Document:
    doc = Document()

    # Título y metadatos
    doc.add_paragraph("Documento de Control de Cambios (DCC)").style = doc.styles['Title']
    add_kv_paragraph(doc, "Proyecto", "Alerta Raven")
    add_kv_paragraph(doc, "Código/ID del Proyecto", "PROYECTO-SW-001")
    add_kv_paragraph(doc, "Versión", "1.0")
    add_kv_paragraph(doc, "Fecha de Emisión", "13/11/2025")
    add_kv_paragraph(doc, "Autor", "Jorge Galindo")
    add_kv_paragraph(doc, "Aprobadores", "Project Manager y Product Owner")

    # 1. Propósito
    add_heading(doc, "1. Propósito", level=1)
    add_bullets(doc, [
        "Define el proceso formal de gestión, control y trazabilidad de cambios del módulo main.py de la API de Alerta Raven.",
        "Minimiza impactos de cambios no planificados y evita desviaciones de alcance.",
        "Garantiza análisis, evaluación, aprobación o rechazo con registro auditable.",
        "Mantiene integridad de líneas base: requisitos, diseño, cronograma y presupuesto.",
    ])

    # 2. Alcance
    add_heading(doc, "2. Alcance", level=1)
    add_bullets(doc, [
        "Cambios que afecten líneas base aprobadas del módulo main.py y sus interfaces.",
        "Endpoints y rutas web: /api/auth/login, /dashboard, /logout, /health y estáticos en /static.",
        "Autenticación y seguridad: JWT en cookies (access_token), create_web_token, verify_web_auth.",
        "Ciclo de vida de la app (lifespan): inicialización DB, heartbeat WebSocket, snapshot de métricas, carga de modelo ML.",
        "Integración con database.py, websocket_manager.py, ml.py y configuración CORS.",
        "Pruebas en tests/*.py y documentación técnica/usuario vinculada.",
        "Fases: Planificación, Diseño, Desarrollo, Implementación y Verificación.",
    ])

    # 3. Proceso de Control de Cambios
    add_heading(doc, "3. Proceso de Control de Cambios", level=1)

    add_heading(doc, "3.1 Iniciación de la Solicitud de Cambio (RSC)", level=2)
    add_bullets(doc, [
        "Cualquier miembro/stakeholder autorizado inicia una RSC con descripción, justificación y alcance en main.py.",
    ])

    add_heading(doc, "3.2 Registro Inicial y Asignación", level=2)
    add_bullets(doc, [
        "Registrar en el Registro de Cambios con ID único.",
        "Asignar responsable para revisión preliminar de completitud.",
    ])

    add_heading(doc, "3.3 Análisis de Impacto Detallado", level=2)
    add_bullets(doc, [
        "Alcance: endpoints y funciones afectadas (autenticación, rutas web, salud, alertas).",
        "Cronograma: estimaciones por iteración/sprint.",
        "Coste/Presupuesto: librerías (p. ej., pyotp), tiempo y pruebas.",
        "Recursos: backend, QA, seguridad; coordinación con frontend si aplica.",
        "Riesgos: regresiones en login/acceso, seguridad, estabilidad de lifespan.",
        "Calidad: cobertura de tests, estándares FastAPI y seguridad.",
        "Técnico: arquitectura de autenticación, cookies, dependencias (requirements.txt).",
    ])

    add_heading(doc, "3.4 Evaluación y Decisión (CCC)", level=2)
    add_bullets(doc, [
        "Decisiones: Aprobado, Rechazado, Diferido o Cancelado.",
        "Registrar decisión, fecha, comentarios y condiciones (p. ej., emitir token solo tras segundo factor).",
    ])

    add_heading(doc, "3.5 Implementación del Cambio Aprobado", level=2)
    add_bullets(doc, [
        "Asignar responsable técnico y plan de actividades.",
        "Modificar main.py manteniendo estilo; actualizar requirements.txt si se añaden dependencias.",
        "Ajustar endpoints/middleware; actualizar static/login.html si el flujo cambia.",
    ])

    add_heading(doc, "3.6 Verificación y Cierre", level=2)
    add_bullets(doc, [
        "Ejecutar tests (tests/test_unit.py, test_integration_api.py, test_websocket.py, test_performance.py).",
        "Verificación manual: login en http://127.0.0.1:8000/login y acceso a /dashboard.",
        "Si es correcto y sin regresiones, marcar RSC como Cerrada en el Registro.",
    ])

    # 4. Líneas Base del Módulo
    add_heading(doc, "4. Líneas Base del Módulo main.py", level=1)
    add_bullets(doc, [
        "Autenticación web: WEB_JWT_SECRET, WEB_JWT_ALGORITHM, ADMIN_USER, ADMIN_PASS.",
        "Funciones: create_web_token(username), verify_web_auth(request).",
        "Endpoints: POST /api/auth/login (cookie access_token httponly), /logout, /dashboard protegido.",
        "Infraestructura: lifespan (DB, heartbeat WebSocket, métricas, ML), estáticos y CORS.",
        "API: /health, alertas y eventos según modelos; PWA (manifest y service worker).",
        "Pruebas: suites en tests/ y BD alertas.db.",
    ])

    # 5. Roles
    add_heading(doc, "5. Roles y Responsabilidades", level=1)
    add_bullets(doc, [
        "Project Manager: prioriza y monitorea cronograma/presupuesto.",
        "Product Owner: define alcance y criterios de aceptación.",
        "Tech Lead Backend: diseña cambios y revisa PRs.",
        "QA Lead: plan y ejecución de pruebas; no regresiones.",
        "Security Officer: valida seguridad (JWT, cookies, MFA).",
        "Desarrollador Backend: implementa cambios y documentación.",
    ])

    # 6. Formulario RSC
    add_heading(doc, "6. Formulario de Solicitud de Cambio (RSC) — Plantilla", level=1)
    add_bullets(doc, [
        "ID de RSC: AR-API-MAIN-RSC-YYYY-XXX",
        "Solicitante, Fecha, Título del cambio, Descripción detallada, Justificación.",
        "Componentes afectados: main.py, static/login.html, requirements.txt, database.py (si aplica), websocket_manager.py (si aplica).",
        "Endpoints afectados, Impacto esperado (alcance, cronograma, coste, recursos, riesgos, calidad, técnico).",
        "Dependencias/Librerías, Pruebas requeridas, Plan de implementación, Plan de rollback.",
        "Aprobación CCC, Comentarios del CCC, Responsable, Fecha de implementación, Estado.",
    ])

    # 7. Registro de Cambios — Tabla
    add_heading(doc, "7. Registro de Cambios — Plantilla", level=1)
    table = doc.add_table(rows=2, cols=8)
    table.style = 'Light Grid'
    headers = [
        "ID RSC", "Fecha", "Título", "Estado",
        "Componente", "Decisión CCC", "Responsable", "Notas"
    ]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    # Fila vacía inicial
    for i in range(8):
        table.cell(1, i).text = ""

    # 8. Criterios de Aprobación
    add_heading(doc, "8. Criterios de Aprobación", level=1)
    add_bullets(doc, [
        "Cumple seguridad (JWT, cookies httponly, samesite adecuado).",
        "Cobertura de pruebas actualizada; suites pasan sin errores.",
        "Sin regresiones en rutas críticas (/dashboard, /login, /api/auth/login, /health).",
        "Documentación y configuración actualizadas.",
        "Validación manual del flujo de login y acceso.",
    ])

    # 9. Versionado y Trazabilidad
    add_heading(doc, "9. Versionado y Trazabilidad", level=1)
    add_bullets(doc, [
        "Versionado semántico cuando cambien contratos de endpoints.",
        "RSC referenciada en commits y PRs; historial del CCC accesible.",
    ])

    # 10. Herramientas y Artefactos
    add_heading(doc, "10. Herramientas y Artefactos", level=1)
    add_bullets(doc, [
        "Código: c:/Users/Alejandro/Desktop/AlertaRaven5/api/main.py",
        "Pruebas: c:/Users/Alejandro/Desktop/AlertaRaven5/api/tests/",
        "Frontend: c:/Users/Alejandro/Desktop/AlertaRaven5/api/static/login.html",
        "Dependencias: c:/Users/Alejandro/Desktop/AlertaRaven5/api/requirements.txt",
        "Base de datos: c:/Users/Alejandro/Desktop/AlertaRaven5/api/alertas.db",
    ])

    # 11. Ejemplo de Cambio (MFA)
    add_heading(doc, "11. Ejemplo de Cambio Relevante (MFA)", level=1)
    add_bullets(doc, [
        "Añadir autenticación multifactor (TOTP) en main.py.",
        "Endpoints: POST /api/auth/mfa/challenge y POST /api/auth/mfa/verify.",
        "Emitir access_token solo tras verificación MFA.",
        "requirements.txt: incluir pyotp; frontend: ajustar static/login.html.",
        "Pruebas unitarias e integración del flujo completo.",
    ])

    # 12. Comunicación
    add_heading(doc, "12. Comunicación", level=1)
    add_bullets(doc, [
        "Comunicar cambios aprobados con resumen, impacto, fecha y acciones de QA.",
        "Actualizar manuales si cambia el flujo de login.",
    ])

    return doc


def main():
    out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.abspath(os.path.join(out_dir, "DCC_Alerta_Raven_mainpy.docx"))
    doc = build_document()
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()